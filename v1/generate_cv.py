import os
import argparse
import tempfile
from datetime import datetime
from pathlib import Path
from markdown2 import markdown
import docx
import openai
from docx import Document
from PyPDF2 import PdfReader

# === Directory Management ===
def create_profile_folder(base_dir, profile_name, version):
    output_dir = Path(base_dir) / profile_name / version
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir

# === File Parsing Utilities ===
def extract_text_from_file(file_path):
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif suffix in [".md", ".txt"]:
        return Path(file_path).read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file format: {suffix}")

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

# === Profile Generation Pipeline ===
class ResumeGenerator:
    def __init__(self, profile_name, base_dir="profiles"):
        self.profile_name = profile_name
        self.base_dir = base_dir

    def generate_version(self):
        timestamp = datetime.now().strftime("%Y%m%d%H%M")
        return f"v{timestamp}"

    def generate_markdown_profile(self, documents, jd_text, template_path, version):
        extracted_info = self.extract_info_from_documents(documents, jd_text)
        markdown_text = self.fill_template_with_info(template_path, extracted_info)

        output_dir = create_profile_folder(self.base_dir, self.profile_name, version)
        md_file_path = output_dir / f"{self.profile_name}.{version}.md"
        with open(md_file_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
        return md_file_path, output_dir

    def generate_outputs(self, markdown_path, output_dir, version):
        with open(markdown_path, "r", encoding="utf-8") as f:
            md_text = f.read()

        doc_path = output_dir / f"{self.profile_name}.{version}.docx"
        self.convert_markdown_to_word(md_text, doc_path)

        html_path = output_dir / f"{self.profile_name}.{version}.html"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(markdown(md_text))

        return doc_path, html_path

    def update_profile(self, new_inputs, old_md_path, new_jd_text):
        name, old_version = self.parse_filename(old_md_path)
        new_version = self.generate_version()
        output_dir = create_profile_folder(self.base_dir, name, new_version)

        updated_info = self.merge_old_and_new_info(old_md_path, new_inputs, new_jd_text)
        new_md_path = output_dir / f"{name}.{new_version}.md"
        with open(new_md_path, "w", encoding="utf-8") as f:
            f.write(updated_info)

        self.generate_outputs(new_md_path, output_dir, new_version)
        return output_dir

    # === AI/NLP Based Methods ===
    def extract_info_from_documents(self, documents, jd_text):
        all_text = "\n\n".join(extract_text_from_file(doc) for doc in documents)
        prompt = f"""
Extract and summarize the most relevant experience from the following user documents to match the job description.

Job Description:
{jd_text}

User Documents:
{all_text}

Format your response as markdown under headings like Experience, Skills, Education, etc.
"""
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response["choices"][0]["message"]["content"]

    def fill_template_with_info(self, template_path, extracted_info):
        return extracted_info

    def convert_markdown_to_word(self, md_text, doc_path):
        doc = docx.Document()
        for line in md_text.splitlines():
            doc.add_paragraph(line)
        doc.save(doc_path)

    def merge_old_and_new_info(self, old_md_path, new_inputs, new_jd_text):
        old_text = Path(old_md_path).read_text(encoding="utf-8")
        new_text = "\n\n".join(extract_text_from_file(doc) for doc in new_inputs)
        prompt = f"""
Update the following markdown resume using these new materials and job description.

Old Resume:
{old_text}

New Inputs:
{new_text}

New Job Description:
{new_jd_text}

Return updated markdown resume.
"""
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response["choices"][0]["message"]["content"]

    def parse_filename(self, md_path):
        parts = Path(md_path).stem.split(".")
        return parts[0], parts[1]

# === CLI Support ===
def main():
    parser = argparse.ArgumentParser(description="Generate structured resumes using AI")
    parser.add_argument("--profile", required=True, help="Profile name")
    parser.add_argument("--docs", nargs="+", required=True, help="List of input files (PDF, DOCX, MD, etc.)")
    parser.add_argument("--jd", required=True, help="Path to job description file")
    parser.add_argument("--template", required=True, help="Path to markdown template file")
    parser.add_argument("--base_dir", default="profiles", help="Base output directory")
    parser.add_argument("--openai_key_env", default="OPENAI_API_KEY", help="Environment variable for OpenAI API key")
    args = parser.parse_args()

    openai.api_key = os.getenv(args.openai_key_env)
    if not openai.api_key:
        raise ValueError("OpenAI API key not found in environment variable")

    jd_text = Path(args.jd).read_text(encoding="utf-8")
    generator = ResumeGenerator(args.profile, args.base_dir)
    version = generator.generate_version()

    md_path, output_dir = generator.generate_markdown_profile(args.docs, jd_text, args.template, version)
    generator.generate_outputs(md_path, output_dir, version)
    print(f"Files saved to: {output_dir}")

if __name__ == "__main__":
    main()
