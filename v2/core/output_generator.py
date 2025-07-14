"""
Output generator for coordinating the generation of different resume formats.
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

try:
    from ..utils import FileUtils, VersionManager
    from ..generators import MarkdownGenerator, WordGenerator, WebsiteGenerator
except ImportError:
    # 回退到绝对导入（当作为独立脚本运行时）
    import sys
    import os
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from utils import FileUtils, VersionManager
    from generators import MarkdownGenerator, WordGenerator, WebsiteGenerator

class OutputGenerator:
    """Coordinates the generation of multiple output formats for resumes."""
    
    def __init__(self, base_output_dir: str = "profiles"):
        """
        Initialize output generator.
        
        Args:
            base_output_dir: Base directory for all profile outputs
        """
        self.base_output_dir = Path(base_output_dir)
        self.logger = logging.getLogger(__name__)
        
        # Initialize utilities and generators
        self.file_utils = FileUtils()
        self.version_manager = VersionManager()
        self.markdown_generator = MarkdownGenerator()
        self.word_generator = WordGenerator()
        self.website_generator = WebsiteGenerator()
        
    def generate_complete_profile(self, profile_data: Dict[str, Any], 
                                profile_name: str, style_analysis: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """
        Generate complete profile with all output formats.
        
        Args:
            profile_data: Profile content and metadata
            profile_name: Name for the profile
            style_analysis: Style analysis from AI processor (optional)
            
        Returns:
            Dictionary mapping format names to output file paths
        """
        try:
            # Generate version and create output directory
            version = self.version_manager.generate_version()
            output_dir = self._create_profile_directory(profile_name, version)
            
            # Store metadata
            metadata = {
                'profile_name': profile_name,
                'version': version,
                'created_at': datetime.now().isoformat(),
                'style_analysis': style_analysis or {},
                'source_documents': profile_data.get('source_documents', [])
            }
            
            output_paths = {}
            
            # Generate markdown file (base format)
            md_path = self.markdown_generator.generate(
                profile_data, output_dir, profile_name, version
            )
            output_paths['markdown'] = str(md_path)
            
            # Generate Word document
            word_path = self.word_generator.generate(
                profile_data, output_dir, profile_name, version, style_analysis
            )
            output_paths['word'] = str(word_path)
            
            # Generate website/HTML
            website_path = self.website_generator.generate(
                profile_data, output_dir, profile_name, version, style_analysis
            )
            output_paths['website'] = str(website_path)
            
            # Save metadata
            metadata_path = output_dir / "metadata.json"
            self.file_utils.save_json(metadata, metadata_path)
            output_paths['metadata'] = str(metadata_path)
            
            self.logger.info(f"Generated complete profile for {profile_name} v{version}")
            return output_paths
            
        except Exception as e:
            self.logger.error(f"Error generating profile: {str(e)}")
            raise Exception(f"Failed to generate complete profile: {str(e)}")
    
    def update_existing_profile(self, old_profile_path: str, updated_content: str, 
                              new_style_analysis: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """
        Update an existing profile with new content.
        
        Args:
            old_profile_path: Path to existing markdown profile
            updated_content: New profile content
            new_style_analysis: Updated style analysis (optional)
            
        Returns:
            Dictionary mapping format names to new output file paths
        """
        try:
            # Parse old profile information
            old_path = Path(old_profile_path)
            profile_name, old_version = self._parse_profile_filename(old_path.name)
            
            # Create new version
            new_version = self.version_manager.generate_version()
            output_dir = self._create_profile_directory(profile_name, new_version)
            
            # Prepare profile data
            profile_data = {
                'content': updated_content,
                'profile_name': profile_name,
                'previous_version': old_version,
                'updated_at': datetime.now().isoformat()
            }
            
            # Generate all formats with new content
            return self.generate_complete_profile(
                profile_data, profile_name, new_style_analysis
            )
            
        except Exception as e:
            self.logger.error(f"Error updating profile: {str(e)}")
            raise Exception(f"Failed to update profile: {str(e)}")
    
    def list_profiles(self) -> List[Dict[str, Any]]:
        """
        List all available profiles with their versions.
        
        Returns:
            List of profile information dictionaries
        """
        profiles = []
        
        if not self.base_output_dir.exists():
            return profiles
            
        try:
            for profile_dir in self.base_output_dir.iterdir():
                if profile_dir.is_dir():
                    profile_info = self._get_profile_info(profile_dir)
                    if profile_info:
                        profiles.append(profile_info)
                        
            return sorted(profiles, key=lambda x: x.get('last_updated', ''))
            
        except Exception as e:
            self.logger.error(f"Error listing profiles: {str(e)}")
            return []
    
    def get_profile_versions(self, profile_name: str) -> List[str]:
        """
        Get all versions for a specific profile.
        
        Args:
            profile_name: Name of the profile
            
        Returns:
            List of version strings, sorted by creation time
        """
        profile_dir = self.base_output_dir / profile_name
        
        if not profile_dir.exists():
            return []
            
        versions = []
        for version_dir in profile_dir.iterdir():
            if version_dir.is_dir():
                versions.append(version_dir.name)
                
        return sorted(versions)
    
    def cleanup_old_versions(self, profile_name: str, keep_versions: int = 5) -> List[str]:
        """
        Clean up old versions of a profile, keeping only the most recent ones.
        
        Args:
            profile_name: Name of the profile
            keep_versions: Number of versions to keep
            
        Returns:
            List of removed version directories
        """
        versions = self.get_profile_versions(profile_name)
        
        if len(versions) <= keep_versions:
            return []
            
        # Remove oldest versions
        versions_to_remove = versions[:-keep_versions]
        removed = []
        
        for version in versions_to_remove:
            version_path = self.base_output_dir / profile_name / version
            try:
                self.file_utils.remove_directory(version_path)
                removed.append(version)
                self.logger.info(f"Removed old version: {profile_name}/{version}")
            except Exception as e:
                self.logger.error(f"Failed to remove {version}: {str(e)}")
                
        return removed
    
    def _create_profile_directory(self, profile_name: str, version: str) -> Path:
        """Create directory structure for profile output."""
        output_dir = self.base_output_dir / profile_name / version
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir
    
    def _parse_profile_filename(self, filename: str) -> tuple:
        """Parse profile filename to extract name and version."""
        # Expected format: {name}.{version}.md
        parts = filename.replace('.md', '').split('.')
        if len(parts) >= 2:
            name = parts[0]
            version = '.'.join(parts[1:])
            return name, version
        else:
            return parts[0], 'unknown'
    
    def _get_profile_info(self, profile_dir: Path) -> Optional[Dict[str, Any]]:
        """Get information about a profile from its directory."""
        try:
            versions = [d.name for d in profile_dir.iterdir() if d.is_dir()]
            if not versions:
                return None
                
            latest_version = sorted(versions)[-1]
            latest_dir = profile_dir / latest_version
            
            # Try to load metadata
            metadata_path = latest_dir / "metadata.json"
            metadata = {}
            if metadata_path.exists():
                metadata = self.file_utils.load_json(metadata_path)
            
            return {
                'name': profile_dir.name,
                'latest_version': latest_version,
                'total_versions': len(versions),
                'last_updated': metadata.get('created_at', 'unknown'),
                'output_dir': str(latest_dir),
                'metadata': metadata
            }
            
        except Exception as e:
            self.logger.error(f"Error getting profile info for {profile_dir}: {str(e)}")
            return None
    
    def generate_word(self, content: str, output_path: str) -> None:
        """
        Generate Word document from markdown content.
        
        Args:
            content: Markdown content to convert
            output_path: Path where to save the Word document
        """
        try:
            # 尝试使用简化的Word生成
            try:
                from docx import Document
                from docx.shared import Inches
                
                # 创建新文档
                doc = Document()
                
                # 设置页边距
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)
                
                # 简单转换markdown到Word
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # 处理标题
                    if line.startswith('# '):
                        p = doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        p = doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        p = doc.add_heading(line[4:], level=3)
                    else:
                        # 普通段落
                        p = doc.add_paragraph(line)
                
                # 保存文档
                doc.save(output_path)
                self.logger.info(f"Word document generated successfully: {output_path}")
                
            except ImportError:
                # python-docx 不可用，创建RTF格式替代
                self.logger.warning("python-docx not available, creating RTF file instead")
                rtf_path = output_path.replace('.docx', '.rtf')
                
                # 简单的RTF格式
                rtf_content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}} "
                rtf_content += content.replace('\n', '\\par ')
                rtf_content += "}"
                
                with open(rtf_path, 'w', encoding='utf-8') as f:
                    f.write(rtf_content)
                    
                self.logger.info(f"RTF document generated successfully: {rtf_path}")
                
        except Exception as e:
            self.logger.error(f"Error generating Word document: {str(e)}")
            # 创建简单的文本文件作为回退
            try:
                txt_path = output_path.replace('.docx', '.txt')
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.logger.info(f"Text file created as fallback: {txt_path}")
            except Exception as fallback_error:
                self.logger.error(f"Fallback text file creation also failed: {fallback_error}")
                raise
    
    def generate_website(self, content: str, output_dir: str, filename_base: str) -> None:
        """
        Generate HTML website from markdown content.
        
        Args:
            content: Markdown content to convert
            output_dir: Directory where to save the HTML files
            filename_base: Base filename for the HTML file
        """
        try:
            output_path = Path(output_dir) / f"{filename_base}.html"
            
            # 尝试使用markdown转HTML
            try:
                import markdown
                html_content = markdown.markdown(content)
            except ImportError:
                # 如果没有markdown库，进行简单转换
                self.logger.warning("markdown library not available, using simple HTML conversion")
                html_content = self._simple_markdown_to_html(content)
            
            # 创建完整的HTML文档
            full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{filename_base} - Resume</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }}
        .container {{
            background-color: white;
            padding: 40px;
            border-radius: 10px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
        }}
        h3 {{
            color: #7f8c8d;
        }}
        p {{
            margin-bottom: 15px;
        }}
        ul {{
            margin-bottom: 15px;
        }}
        @media print {{
            body {{
                background-color: white;
            }}
            .container {{
                box-shadow: none;
                padding: 0;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        {html_content}
    </div>
</body>
</html>"""
            
            # 保存HTML文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
                
            self.logger.info(f"HTML website generated successfully: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Error generating HTML website: {str(e)}")
            raise
    
    def _simple_markdown_to_html(self, markdown_content: str) -> str:
        """
        Simple markdown to HTML conversion when markdown library is not available.
        
        Args:
            markdown_content: Markdown content to convert
            
        Returns:
            HTML content
        """
        lines = markdown_content.split('\n')
        html_lines = []
        in_list = False
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append('<br>')
                continue
            
            # 处理标题
            if line.startswith('### '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('## '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('# '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h1>{line[2:]}</h1>')
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                if not in_list:
                    html_lines.append('<ul>')
                    in_list = True
                html_lines.append(f'<li>{line[2:]}</li>')
            # 普通段落
            else:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<p>{line}</p>')
        
        # 关闭未关闭的列表
        if in_list:
            html_lines.append('</ul>')
        
        return '\n'.join(html_lines) 