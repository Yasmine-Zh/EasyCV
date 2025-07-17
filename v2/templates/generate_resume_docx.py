# Re-run necessary imports and code due to state reset
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Define directory for DOCX templates
docx_template_dir = Path("templates/resume_docx")
docx_template_dir.mkdir(parents=True, exist_ok=True)

def create_docx_template(filename: str, style: str = "classic") -> Path:
    doc = Document()

    # Title
    title = doc.add_heading("{{ name }}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("{{ location }} · {{ email }} · {{ phone }} · {{ linkedin }}")
    run.font.size = Pt(10)

    doc.add_paragraph("\n")

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("{{ summary }}")

    # Experience
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("{% for exp in experiences %}")
    doc.add_paragraph("{{ exp.title }} at {{ exp.organization }} ({{ exp.time }})", style="List Bullet")
    doc.add_paragraph("{% for line in exp.content.split('\\n') %}- {{ line }}{% endfor %}", style="Normal")
    doc.add_paragraph("{% endfor %}")

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("{% for edu in education %}")
    doc.add_paragraph("{{ edu.degree }} at {{ edu.school }} ({{ edu.time }})", style="List Bullet")
    doc.add_paragraph("{{ edu.content }}", style="Normal")
    doc.add_paragraph("{% endfor %}")

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("{% for skill in skills %}• {{ skill.skill }} ({{ skill.degree }}){% endfor %}")

    # Achievements
    doc.add_heading("Achievements", level=1)
    doc.add_paragraph("{% for ach in achievements %}• {{ ach }}{% endfor %}")

    path = docx_template_dir / filename
    doc.save(path)
    return path

# Create two templates with different intended styles
template1 = create_docx_template("classic_template.docx", style="classic")
template2 = create_docx_template("modern_template.docx", style="modern")

[template1.name, template2.name]
