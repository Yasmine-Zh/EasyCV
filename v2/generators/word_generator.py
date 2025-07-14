"""
Word document generator for resume profiles.
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None

class WordGenerator:
    """Generator for Word document format resume profiles."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        if Document is None:
            self.logger.warning("python-docx not available. Word generation will be limited.")
    
    def generate(self, profile_data: Dict[str, Any], output_dir: Path, 
                profile_name: str, version: str, style_analysis: Optional[Dict[str, Any]] = None) -> Path:
        """
        Generate Word document resume file.
        
        Args:
            profile_data: Profile content and metadata
            output_dir: Directory for output files
            profile_name: Name of the profile
            version: Version string
            style_analysis: Style analysis for formatting
            
        Returns:
            Path to generated Word document
        """
        if Document is None:
            raise ImportError("python-docx is required for Word document generation")
            
        try:
            # Generate filename
            filename = f"{profile_name}.{version}.docx"
            output_path = output_dir / filename
            
            # Get content from profile data
            content = profile_data.get('content', '')
            
            # Create document
            doc = Document()
            
            # Apply styling
            self._setup_document_styles(doc, style_analysis)
            
            # Convert markdown content to Word format
            self._convert_markdown_to_word(doc, content)
            
            # Save document
            doc.save(output_path)
            
            self.logger.info(f"Generated Word document: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating Word document: {str(e)}")
            raise Exception(f"Failed to generate Word document: {str(e)}")
    
    def _setup_document_styles(self, doc: Document, style_analysis: Optional[Dict[str, Any]] = None):
        """Setup document styles based on analysis."""
        try:
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Create heading styles if they don't exist
            self._create_heading_styles(doc)
            
            # Apply custom styling based on analysis
            if style_analysis:
                self._apply_style_analysis(doc, style_analysis)
                
        except Exception as e:
            self.logger.warning(f"Error setting up styles: {str(e)}")
    
    def _create_heading_styles(self, doc: Document):
        """Create custom heading styles."""
        try:
            # Create or modify heading styles
            for level in range(1, 4):
                style_name = f'Heading {level}'
                if style_name in doc.styles:
                    style = doc.styles[style_name]
                else:
                    style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                
                # Configure heading style
                font = style.font
                font.bold = True
                font.name = 'Calibri'
                
                if level == 1:
                    font.size = Pt(16)
                elif level == 2:
                    font.size = Pt(14)
                else:
                    font.size = Pt(12)
                    
        except Exception as e:
            self.logger.warning(f"Error creating heading styles: {str(e)}")
    
    def _apply_style_analysis(self, doc: Document, style_analysis: Dict[str, Any]):
        """Apply styling based on analysis."""
        # This could be expanded to apply specific formatting
        # based on the AI style analysis
        pass
    
    def _convert_markdown_to_word(self, doc: Document, markdown_content: str):
        """Convert markdown content to Word document format."""
        try:
            # Remove YAML front matter if present
            content = self._remove_yaml_frontmatter(markdown_content)
            
            # Split content into lines
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    # Add empty paragraph for spacing
                    doc.add_paragraph()
                    continue
                
                # Handle different markdown elements
                if line.startswith('# '):
                    # Heading 1
                    heading = doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # Heading 2
                    heading = doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # Heading 3
                    heading = doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or re.match(r'^\d+\. ', line):
                    # Numbered list
                    text = re.sub(r'^\d+\. ', '', line)
                    p = doc.add_paragraph(text, style='List Number')
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text (simple case)
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                elif line.startswith('*') and line.endswith('*'):
                    # Italic text (simple case)
                    p = doc.add_paragraph()
                    run = p.add_run(line[1:-1])
                    run.italic = True
                else:
                    # Regular paragraph
                    self._add_formatted_paragraph(doc, line)
                    
        except Exception as e:
            self.logger.error(f"Error converting markdown: {str(e)}")
            # Fallback: add as plain text
            doc.add_paragraph(markdown_content)
    
    def _add_formatted_paragraph(self, doc: Document, text: str):
        """Add paragraph with inline formatting."""
        try:
            p = doc.add_paragraph()
            
            # Simple inline formatting (can be enhanced)
            parts = self._parse_inline_formatting(text)
            
            for part_text, formatting in parts:
                run = p.add_run(part_text)
                
                if 'bold' in formatting:
                    run.bold = True
                if 'italic' in formatting:
                    run.italic = True
                    
        except Exception as e:
            self.logger.warning(f"Error formatting paragraph: {str(e)}")
            # Fallback
            doc.add_paragraph(text)
    
    def _parse_inline_formatting(self, text: str) -> list:
        """Parse inline markdown formatting."""
        # Simple implementation - can be enhanced
        # Returns list of (text, formatting) tuples
        
        # For now, just return text without formatting
        return [(text, [])]
    
    def _remove_yaml_frontmatter(self, content: str) -> str:
        """Remove YAML front matter from content."""
        if not content.startswith('---'):
            return content
            
        lines = content.split('\n')
        yaml_end = -1
        
        for i, line in enumerate(lines[1:], 1):
            if line.strip() == '---':
                yaml_end = i
                break
        
        if yaml_end > 0:
            return '\n'.join(lines[yaml_end + 1:]).lstrip('\n')
        else:
            return content
    
    def add_header_footer(self, doc: Document, profile_name: str, version: str):
        """Add header and footer to document."""
        try:
            # Add header
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"{profile_name} Resume"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add footer
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Generated by EasyCV - Version {version}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            self.logger.warning(f"Error adding header/footer: {str(e)}")
    
    def set_page_margins(self, doc: Document, top: float = 1.0, bottom: float = 1.0, 
                        left: float = 1.0, right: float = 1.0):
        """Set page margins in inches."""
        try:
            section = doc.sections[0]
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
            
        except Exception as e:
            self.logger.warning(f"Error setting margins: {str(e)}")
    
    def validate_document(self, doc_path: Path) -> Dict[str, Any]:
        """Validate generated Word document."""
        try:
            # Check if file exists and is readable
            if not doc_path.exists():
                return {'valid': False, 'error': 'File does not exist'}
                
            file_size = doc_path.stat().st_size
            
            # Try to open document
            if Document is not None:
                try:
                    doc = Document(doc_path)
                    paragraph_count = len(doc.paragraphs)
                    
                    return {
                        'valid': True,
                        'file_size': file_size,
                        'paragraph_count': paragraph_count,
                        'readable': True
                    }
                except Exception as e:
                    return {
                        'valid': False,
                        'file_size': file_size,
                        'error': f"Document unreadable: {str(e)}"
                    }
            else:
                return {
                    'valid': True,
                    'file_size': file_size,
                    'note': 'python-docx not available for validation'
                }
                
        except Exception as e:
            return {'valid': False, 'error': str(e)} 